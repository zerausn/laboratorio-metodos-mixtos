from __future__ import annotations

import argparse
import json
import math
import os
import re
import statistics
import sys
import unicodedata
from collections import defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

import cv2
import fitz
import numpy as np
import pandas as pd
import pytesseract
from docx import Document as DocxDocument
from google.cloud.documentai_toolbox.wrappers.document import Document as WrappedDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from project_paths import DOCAI_JSON_DIR, OUTPUTS_DIR, PDF_PATH, TESSDATA_DIR


MARKER_RE = re.compile(r"(?:pagina\s+)?(\d{1,2})\s+de\s+(\d{1,2})", re.IGNORECASE)
KNOWN_PAGE_OFFSETS = {32: 2, 13: 35, 16: 48, 2: 64, 1: 67}
KEYWORD_LINES = [
    "ingresos",
    "gastos",
    "presupuesto",
    "recaudo",
    "pagos",
    "funcionamiento",
    "inversion",
    "deuda",
    "sobretasa",
    "predial",
    "industria y comercio",
    "resumen",
    "total",
    "consolidado",
]
DEFAULT_TESSERACT = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")


@dataclass
class PdfPageOcr:
    absolute_page: int
    best_rotation: int
    avg_confidence: float
    word_count: int
    score: float
    marker_page: int | None
    marker_total: int | None
    section: str
    suspicious_ratio: float
    difficult: bool
    text: str
    preview: str


@dataclass
class DocAiPage:
    source_file: str
    local_page: int
    marker_page: int | None
    marker_total: int | None
    mapped_pdf_page: int | None
    section: str
    avg_token_confidence: float
    table_count: int
    text_length: int
    preview: str
    text: str


@dataclass
class DocAiTableRecord:
    source_file: str
    local_page: int
    mapped_pdf_page: int | None
    section: str
    table_index: int
    rows: int
    cols: int
    sheet_name: str
    csv_path: str


def strip_accents(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))


def normalize_text(text: str) -> str:
    base = strip_accents(text or "")
    return re.sub(r"\s+", " ", base).strip().lower()


def preview_text(text: str, limit: int = 180) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    return compact[:limit]


def safe_sheet_name(name: str) -> str:
    clean = re.sub(r"[\[\]\*:/\\?]", "_", name)
    return clean[:31] or "sheet"


def extract_marker(text: str) -> tuple[int | None, int | None]:
    matches: list[tuple[int, int]] = []
    for match in MARKER_RE.finditer(strip_accents(text or "")):
        page_num = int(match.group(1))
        total_num = int(match.group(2))
        if 1 <= page_num <= total_num <= 100:
            matches.append((page_num, total_num))
    if not matches:
        return None, None
    return matches[-1]


def detect_section(text: str) -> str:
    normalized = normalize_text(text)
    if "resumen ejecucion presupuestal" in normalized:
        return "resumen"
    if "gastos" in normalized or "ejecucion presupuestal de gastos" in normalized:
        return "gastos"
    if "ingresos" in normalized or "ejecucion de ingresos" in normalized:
        return "ingresos"
    return "sin_clasificar"


def section_from_marker_total(marker_total: int | None) -> str:
    if marker_total == 32:
        return "ingresos"
    if marker_total in {13, 16}:
        return "gastos"
    if marker_total == 2:
        return "resumen_gastos"
    if marker_total == 1:
        return "deuda"
    return "sin_clasificar"


def suspicious_ratio(words: list[str]) -> float:
    if not words:
        return 1.0
    weird = 0
    for word in words:
        cleaned = word.strip()
        if not cleaned:
            continue
        if len(cleaned) == 1 and not cleaned.isdigit():
            weird += 1
            continue
        if re.search(r"[^0-9A-Za-z.,/%$()\-+]", cleaned):
            weird += 1
    return weird / max(len(words), 1)


def mean(values: list[float]) -> float:
    return statistics.fmean(values) if values else 0.0


def flatten_columns(frame: pd.DataFrame) -> pd.DataFrame:
    copy = frame.copy()
    if isinstance(copy.columns, pd.MultiIndex):
        copy.columns = [
            " | ".join(str(part).strip() for part in col if str(part).strip())
            for col in copy.columns.to_list()
        ]
    else:
        copy.columns = [str(col) for col in copy.columns]
    return copy


def table_to_dataframe(table: Any) -> pd.DataFrame:
    body_rows = [list(row) for row in table.body_rows]
    header_rows = [list(row) for row in table.header_rows]
    width = 0
    for row in header_rows + body_rows:
        width = max(width, len(row))
    if width == 0:
        return pd.DataFrame()

    def pad(row: list[str]) -> list[str]:
        padded = list(row)
        while len(padded) < width:
            padded.append("")
        return padded[:width]

    padded_headers = [pad(row) for row in header_rows]
    padded_body = [pad(row) for row in body_rows]

    if padded_headers:
        columns = []
        for col_index in range(width):
            parts = [str(row[col_index]).strip() for row in padded_headers if str(row[col_index]).strip()]
            columns.append(" | ".join(parts) if parts else f"col_{col_index + 1}")
    else:
        columns = [f"col_{col_index + 1}" for col_index in range(width)]

    return pd.DataFrame(padded_body, columns=columns)


def to_records_dataframe(records: list[dict[str, Any]]) -> pd.DataFrame:
    return pd.DataFrame(records) if records else pd.DataFrame()


def configure_tesseract(tesseract_cmd: Path | None, tessdata_dir: Path | None) -> None:
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = str(tesseract_cmd)
    if tessdata_dir:
        os.environ["TESSDATA_PREFIX"] = str(tessdata_dir)


def resolve_tesseract(cmd_arg: str | None) -> Path | None:
    if cmd_arg:
        path = Path(cmd_arg)
        return path if path.exists() else None
    return DEFAULT_TESSERACT if DEFAULT_TESSERACT.exists() else None


def render_page(page: fitz.Page, zoom: float) -> np.ndarray:
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    image = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        image = cv2.cvtColor(image, cv2.COLOR_BGRA2BGR)
    return image


def preprocess_for_ocr(image: np.ndarray) -> np.ndarray:
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    threshold = cv2.adaptiveThreshold(
        denoised,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        11,
    )
    return threshold


def rotate_image(image: np.ndarray, rotation: int) -> np.ndarray:
    if rotation == 90:
        return cv2.rotate(image, cv2.ROTATE_90_CLOCKWISE)
    if rotation == 180:
        return cv2.rotate(image, cv2.ROTATE_180)
    if rotation == 270:
        return cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
    return image


def rebuild_text_from_data(data: dict[str, list[Any]]) -> str:
    chunks: list[str] = []
    last_signature: tuple[int, int, int] | None = None
    for idx, raw_word in enumerate(data["text"]):
        word = str(raw_word).strip()
        if not word:
            continue
        signature = (
            int(data["block_num"][idx]),
            int(data["par_num"][idx]),
            int(data["line_num"][idx]),
        )
        if last_signature is None:
            chunks.append(word)
        elif signature == last_signature:
            chunks.append(f" {word}")
        else:
            chunks.append(f"\n{word}")
        last_signature = signature
    return "".join(chunks).strip()


def score_ocr(data: dict[str, list[Any]]) -> tuple[float, float, int, float]:
    words = [str(word).strip() for word in data["text"] if str(word).strip()]
    confidences = [
        float(conf)
        for conf in data["conf"]
        if str(conf).strip() not in {"", "-1"} and float(conf) >= 0
    ]
    avg_conf = mean(confidences)
    weird_ratio = suspicious_ratio(words)
    word_count = len(words)
    score = avg_conf * math.log(word_count + 1) * max(0.1, 1 - weird_ratio)
    return avg_conf, weird_ratio, word_count, score


def ocr_pdf_page(page: fitz.Page) -> tuple[PdfPageOcr, list[dict[str, Any]]]:
    base_image = render_page(page, zoom=2.0)
    candidates: list[tuple[int, dict[str, list[Any]], str, float, float, int, float]] = []
    for rotation in (0, 90):
        rotated = rotate_image(base_image, rotation)
        processed = preprocess_for_ocr(rotated)
        data = pytesseract.image_to_data(
            processed,
            lang="spa+eng",
            config="--psm 6",
            output_type=pytesseract.Output.DICT,
        )
        text = rebuild_text_from_data(data)
        avg_conf, weird_ratio, word_count, score = score_ocr(data)
        candidates.append((rotation, data, text, avg_conf, weird_ratio, word_count, score))

    best_rotation, data, text, avg_conf, weird_ratio, word_count, score = max(
        candidates,
        key=lambda item: item[6],
    )
    marker_page, marker_total = extract_marker(text)
    section = detect_section(text)
    if section == "sin_clasificar":
        section = section_from_marker_total(marker_total)
    difficult = avg_conf < 55 or weird_ratio > 0.35 or word_count < 20

    token_rows: list[dict[str, Any]] = []
    for idx, word in enumerate(data["text"]):
        clean = str(word).strip()
        if not clean:
            continue
        conf = float(data["conf"][idx]) if str(data["conf"][idx]).strip() not in {"", "-1"} else -1.0
        token_rows.append(
            {
                "absolute_page": page.number + 1,
                "rotation": best_rotation,
                "word": clean,
                "confidence": conf,
                "left": int(data["left"][idx]),
                "top": int(data["top"][idx]),
                "width": int(data["width"][idx]),
                "height": int(data["height"][idx]),
                "block_num": int(data["block_num"][idx]),
                "par_num": int(data["par_num"][idx]),
                "line_num": int(data["line_num"][idx]),
                "word_num": int(data["word_num"][idx]),
            }
        )

    result = PdfPageOcr(
        absolute_page=page.number + 1,
        best_rotation=best_rotation,
        avg_confidence=round(avg_conf, 2),
        word_count=word_count,
        score=round(score, 2),
        marker_page=marker_page,
        marker_total=marker_total,
        section=section,
        suspicious_ratio=round(weird_ratio, 4),
        difficult=difficult,
        text=text,
        preview=preview_text(text),
    )
    return result, token_rows


def load_docai_pages(docai_dir: Path, export_dir: Path) -> tuple[list[DocAiPage], list[DocAiTableRecord]]:
    pages_out: list[DocAiPage] = []
    tables_out: list[DocAiTableRecord] = []
    hocr_dir = export_dir / "hocr"
    table_csv_dir = export_dir / "tables_csv"
    hocr_dir.mkdir(parents=True, exist_ok=True)
    table_csv_dir.mkdir(parents=True, exist_ok=True)

    for json_path in sorted(docai_dir.glob("document*.json")):
        wrapped = WrappedDocument.from_document_path(str(json_path))
        hocr_path = hocr_dir / f"{json_path.stem}.hocr.html"
        hocr_path.write_text(wrapped.export_hocr_str(title=json_path.stem), encoding="utf-8")

        for page_index, page in enumerate(wrapped.pages, start=1):
            text = page.text or ""
            marker_page, marker_total = extract_marker(text)
            token_confidences = [
                float(token.documentai_object.layout.confidence)
                for token in page.tokens
                if getattr(token.documentai_object, "layout", None) is not None
            ]
            entry = DocAiPage(
                source_file=json_path.name,
                local_page=page_index,
                marker_page=marker_page,
                marker_total=marker_total,
                mapped_pdf_page=None,
                section=detect_section(text),
                avg_token_confidence=round(mean(token_confidences), 4),
                table_count=len(page.tables),
                text_length=len(text),
                preview=preview_text(text),
                text=text,
            )
            pages_out.append(entry)

            for table_index, table in enumerate(page.tables, start=1):
                frame = flatten_columns(table_to_dataframe(table))
                csv_name = safe_sheet_name(f"{json_path.stem}_p{page_index:02d}_t{table_index:02d}") + ".csv"
                csv_path = table_csv_dir / csv_name
                frame.to_csv(csv_path, index=False, encoding="utf-8-sig")
                tables_out.append(
                    DocAiTableRecord(
                        source_file=json_path.name,
                        local_page=page_index,
                        mapped_pdf_page=None,
                        section=entry.section,
                        table_index=table_index,
                        rows=int(frame.shape[0]),
                        cols=int(frame.shape[1]),
                        sheet_name=safe_sheet_name(f"{json_path.stem[:10]}_p{page_index:02d}_t{table_index:02d}"),
                        csv_path=str(csv_path),
                    )
                )

    return pages_out, tables_out


def build_pdf_marker_index(pdf_pages: list[PdfPageOcr]) -> tuple[dict[tuple[int, int], list[int]], dict[int, int]]:
    by_marker: dict[tuple[int, int], list[int]] = defaultdict(list)
    offsets: dict[int, int] = {}
    for item in pdf_pages:
        if item.marker_page is not None and item.marker_total is not None:
            pair = (item.marker_page, item.marker_total)
            by_marker[pair].append(item.absolute_page)
            if item.marker_page == 1 and item.marker_total not in offsets:
                offsets[item.marker_total] = item.absolute_page - 1
    return by_marker, offsets


def map_docai_pages(docai_pages: list[DocAiPage], pdf_pages: list[PdfPageOcr]) -> None:
    marker_index, offsets = build_pdf_marker_index(pdf_pages)
    used_pages: set[int] = set()

    for item in docai_pages:
        if item.marker_page is None or item.marker_total is None:
            continue
        pair = (item.marker_page, item.marker_total)
        candidates = [page for page in marker_index.get(pair, []) if page not in used_pages]
        if len(candidates) == 1:
            item.mapped_pdf_page = candidates[0]
            used_pages.add(candidates[0])
            continue
        offset = offsets.get(item.marker_total)
        if offset is not None:
            candidate = offset + item.marker_page
            if candidate not in used_pages:
                item.mapped_pdf_page = candidate
                used_pages.add(candidate)

    grouped: dict[str, list[DocAiPage]] = defaultdict(list)
    for item in docai_pages:
        grouped[item.source_file].append(item)

    for group in grouped.values():
        group.sort(key=lambda row: row.local_page)
        for idx, row in enumerate(group):
            if row.mapped_pdf_page is not None:
                continue
            prev_page = group[idx - 1].mapped_pdf_page if idx > 0 else None
            next_page = group[idx + 1].mapped_pdf_page if idx + 1 < len(group) else None
            if prev_page is not None and next_page is not None and next_page - prev_page == 2:
                row.mapped_pdf_page = prev_page + 1


def seed_docai_mappings(docai_pages: list[DocAiPage]) -> None:
    for item in docai_pages:
        if item.marker_page is None or item.marker_total is None:
            continue
        offset = KNOWN_PAGE_OFFSETS.get(item.marker_total)
        if offset is not None:
            item.mapped_pdf_page = offset + item.marker_page
        if item.section == "sin_clasificar":
            item.section = section_from_marker_total(item.marker_total)

    grouped: dict[str, list[DocAiPage]] = defaultdict(list)
    for item in docai_pages:
        grouped[item.source_file].append(item)

    for group in grouped.values():
        group.sort(key=lambda row: row.local_page)
        for idx, row in enumerate(group):
            if row.mapped_pdf_page is not None:
                continue
            prev_page = group[idx - 1].mapped_pdf_page if idx > 0 else None
            next_page = group[idx + 1].mapped_pdf_page if idx + 1 < len(group) else None
            if prev_page is not None and next_page is not None and next_page - prev_page == 2:
                row.mapped_pdf_page = prev_page + 1
                continue
            if prev_page is not None:
                row.mapped_pdf_page = prev_page + 1


def synthesize_pdf_pages(
    page_count: int,
    docai_pages: list[DocAiPage],
    ocr_pages: list[PdfPageOcr],
) -> list[PdfPageOcr]:
    merged: dict[int, PdfPageOcr] = {page.absolute_page: page for page in ocr_pages}

    for item in docai_pages:
        if item.mapped_pdf_page is None:
            continue
        words = [word for word in re.split(r"\s+", item.text) if word]
        weird_ratio = suspicious_ratio(words)
        avg_conf = round(item.avg_token_confidence * 100, 2)
        synthetic = PdfPageOcr(
            absolute_page=item.mapped_pdf_page,
            best_rotation=-1,
            avg_confidence=avg_conf,
            word_count=len(words),
            score=avg_conf,
            marker_page=item.marker_page,
            marker_total=item.marker_total,
            section=item.section if item.section != "sin_clasificar" else section_from_marker_total(item.marker_total),
            suspicious_ratio=round(weird_ratio, 4),
            difficult=avg_conf < 80 or weird_ratio > 0.25,
            text=item.text,
            preview=item.preview,
        )
        current = merged.get(item.mapped_pdf_page)
        if current is None or current.best_rotation != -1:
            merged[item.mapped_pdf_page] = synthetic

    return [merged[number] for number in sorted(merged) if 1 <= number <= page_count]


def update_table_mappings(docai_pages: list[DocAiPage], tables: list[DocAiTableRecord]) -> None:
    mapping = {
        (page.source_file, page.local_page): page.mapped_pdf_page
        for page in docai_pages
    }
    for table in tables:
        table.mapped_pdf_page = mapping.get((table.source_file, table.local_page))


def extract_keyword_lines(text: str, max_lines: int = 8) -> list[str]:
    lines = [line.strip() for line in (text or "").splitlines()]
    selected: list[str] = []
    for line in lines:
        lowered = normalize_text(line)
        if not lowered:
            continue
        if any(keyword in lowered for keyword in KEYWORD_LINES) and re.search(r"\d", line):
            selected.append(line)
        if len(selected) >= max_lines:
            break
    return selected


def write_excel(
    out_path: Path,
    pdf_pages: list[PdfPageOcr],
    docai_pages: list[DocAiPage],
    docai_tables: list[DocAiTableRecord],
    table_csv_dir: Path,
) -> None:
    page_inventory = []
    for page in pdf_pages:
        page_inventory.append(
            {
                "absolute_page": page.absolute_page,
                "best_rotation": page.best_rotation,
                "ocr_avg_confidence": page.avg_confidence,
                "ocr_word_count": page.word_count,
                "ocr_score": page.score,
                "marker_page": page.marker_page,
                "marker_total": page.marker_total,
                "section": page.section,
                "suspicious_ratio": page.suspicious_ratio,
                "difficult": page.difficult,
                "preview": page.preview,
            }
        )

    docai_inventory = []
    for page in docai_pages:
        docai_inventory.append(
            {
                "source_file": page.source_file,
                "local_page": page.local_page,
                "marker_page": page.marker_page,
                "marker_total": page.marker_total,
                "mapped_pdf_page": page.mapped_pdf_page,
                "section": page.section,
                "avg_token_confidence": page.avg_token_confidence,
                "table_count": page.table_count,
                "text_length": page.text_length,
                "preview": page.preview,
            }
        )

    tables_index = [asdict(item) for item in docai_tables]

    difficult_pages = [
        {
            "absolute_page": item.absolute_page,
            "ocr_avg_confidence": item.avg_confidence,
            "suspicious_ratio": item.suspicious_ratio,
            "marker_page": item.marker_page,
            "marker_total": item.marker_total,
            "section": item.section,
            "preview": item.preview,
        }
        for item in pdf_pages
        if item.difficult
    ]

    keyword_rows = []
    for item in pdf_pages:
        for line in extract_keyword_lines(item.text):
            keyword_rows.append(
                {
                    "absolute_page": item.absolute_page,
                    "section": item.section,
                    "line": line,
                }
            )

    with pd.ExcelWriter(out_path, engine="xlsxwriter") as writer:
        to_records_dataframe(page_inventory).to_excel(writer, sheet_name="page_inventory", index=False)
        to_records_dataframe(docai_inventory).to_excel(writer, sheet_name="docai_pages", index=False)
        to_records_dataframe(tables_index).to_excel(writer, sheet_name="docai_tables", index=False)
        to_records_dataframe(difficult_pages).to_excel(writer, sheet_name="difficult_pages", index=False)
        to_records_dataframe(keyword_rows).to_excel(writer, sheet_name="keyword_lines", index=False)

        for table in docai_tables:
            csv_path = Path(table.csv_path)
            if not csv_path.exists():
                continue
            frame = pd.read_csv(csv_path)
            frame.to_excel(writer, sheet_name=table.sheet_name, index=False)


def write_text_exports(out_dir: Path, pdf_pages: list[PdfPageOcr], docai_pages: list[DocAiPage]) -> None:
    pdf_text_dir = out_dir / "pdf_ocr_text"
    docai_text_dir = out_dir / "docai_text"
    pdf_text_dir.mkdir(parents=True, exist_ok=True)
    docai_text_dir.mkdir(parents=True, exist_ok=True)

    for page in pdf_pages:
        path = pdf_text_dir / f"page_{page.absolute_page:03d}.txt"
        path.write_text(page.text, encoding="utf-8")

    for page in docai_pages:
        mapped = f"_pdf_{page.mapped_pdf_page:03d}" if page.mapped_pdf_page else ""
        path = docai_text_dir / f"{Path(page.source_file).stem}_page_{page.local_page:02d}{mapped}.txt"
        path.write_text(page.text, encoding="utf-8")


def write_csv_exports(out_dir: Path, page_inventory: list[PdfPageOcr], token_rows: list[dict[str, Any]], docai_pages: list[DocAiPage]) -> None:
    page_frame = pd.DataFrame([asdict(item) for item in page_inventory])
    docai_frame = pd.DataFrame([asdict(item) for item in docai_pages])
    tokens_frame = pd.DataFrame(token_rows)
    page_frame.to_csv(out_dir / "pdf_page_inventory.csv", index=False, encoding="utf-8-sig")
    docai_frame.to_csv(out_dir / "docai_page_inventory.csv", index=False, encoding="utf-8-sig")
    tokens_frame.to_csv(out_dir / "ocr_tokens.csv", index=False, encoding="utf-8-sig")


def build_summary(pdf_pages: list[PdfPageOcr], docai_pages: list[DocAiPage], docai_tables: list[DocAiTableRecord]) -> list[str]:
    total_pages = len(pdf_pages)
    difficult = [page.absolute_page for page in pdf_pages if page.difficult]
    docai_mapped = sum(1 for page in docai_pages if page.mapped_pdf_page is not None)
    section_counts: dict[str, int] = defaultdict(int)
    for item in pdf_pages:
        section_counts[item.section] += 1

    summary_lines = [
        f"PDF total pages: {total_pages}",
        f"Document AI pages loaded: {len(docai_pages)}",
        f"Document AI pages mapped to PDF: {docai_mapped}",
        f"Document AI extracted tables: {len(docai_tables)}",
        f"Difficult PDF pages flagged: {len(difficult)}",
        f"Sections detected: {dict(section_counts)}",
    ]
    if difficult:
        summary_lines.append(
            "Most difficult pages: " + ", ".join(str(page) for page in difficult[:20])
        )
    return summary_lines


def write_pdf_report(
    report_path: Path,
    pdf_pages: list[PdfPageOcr],
    docai_pages: list[DocAiPage],
    docai_tables: list[DocAiTableRecord],
    pdf_path: Path,
    docai_dir: Path,
) -> None:
    styles = getSampleStyleSheet()
    story: list[Any] = []

    story.append(Paragraph("Document exploitation report", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"PDF source: {pdf_path}", styles["Normal"]))
    story.append(Paragraph(f"Document AI directory: {docai_dir}", styles["Normal"]))
    story.append(Spacer(1, 12))

    for line in build_summary(pdf_pages, docai_pages, docai_tables):
        story.append(Paragraph(line, styles["BodyText"]))
        story.append(Spacer(1, 6))

    story.append(Spacer(1, 12))
    story.append(Paragraph("Important data zones", styles["Heading2"]))
    important_rows = [
        ["Area", "Pages", "Why it matters"],
        ["Ingresos", "3-34", "Tax revenues, recaudation, and execution percentages."],
        ["Gastos cover", "35", "Section split between ingresos and gastos."],
        ["Gastos block 1", "36-48", "Execution by spending items extracted from Document AI JSON."],
        ["Gastos block 2", "49-64", "OCR-only recovery area with table-heavy pages."],
        ["Executive summary", "69", "Total, funcionamiento, deuda, and inversion summary."],
    ]
    important_table = Table(important_rows, colWidths=[110, 80, 300])
    important_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(important_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Difficult pages", styles["Heading2"]))
    hard_rows = [["PDF page", "OCR confidence", "Suspicious ratio", "Preview"]]
    for item in [page for page in pdf_pages if page.difficult][:15]:
        hard_rows.append(
            [
                str(item.absolute_page),
                str(item.avg_confidence),
                str(item.suspicious_ratio),
                item.preview[:90],
            ]
        )
    hard_table = Table(hard_rows, colWidths=[60, 80, 90, 310])
    hard_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(hard_table)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Recovered macro figures", styles["Heading2"]))
    for page in pdf_pages:
        if page.absolute_page in {3, 35, 69}:
            story.append(Paragraph(f"PDF page {page.absolute_page}", styles["Heading3"]))
            for line in extract_keyword_lines(page.text):
                story.append(Paragraph(line, styles["BodyText"]))
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(str(report_path), pagesize=letter)
    doc.build(story)


def write_docx_report(
    report_path: Path,
    pdf_pages: list[PdfPageOcr],
    docai_pages: list[DocAiPage],
    docai_tables: list[DocAiTableRecord],
    pdf_path: Path,
    docai_dir: Path,
) -> None:
    doc = DocxDocument()
    doc.add_heading("Document exploitation report", 0)
    doc.add_paragraph(f"PDF source: {pdf_path}")
    doc.add_paragraph(f"Document AI directory: {docai_dir}")

    doc.add_heading("Summary", level=1)
    for line in build_summary(pdf_pages, docai_pages, docai_tables):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Important data zones", level=1)
    doc.add_paragraph("Ingresos: PDF pages 3-34, budget and recaudation tables.")
    doc.add_paragraph("Gastos cover: PDF page 35.")
    doc.add_paragraph("Gastos block 1: PDF pages 36-48, mapped from Document AI.")
    doc.add_paragraph("Gastos block 2: PDF pages 49-64, recovered with local OCR.")
    doc.add_paragraph("Executive summary: PDF page 69.")

    doc.add_heading("Difficult pages", level=1)
    for item in [page for page in pdf_pages if page.difficult][:20]:
        doc.add_paragraph(
            f"PDF page {item.absolute_page}: OCR confidence={item.avg_confidence}, suspicious_ratio={item.suspicious_ratio}, preview={item.preview}"
        )

    doc.save(str(report_path))


def write_processing_manifest(
    manifest_path: Path,
    pdf_pages: list[PdfPageOcr],
    docai_pages: list[DocAiPage],
    docai_tables: list[DocAiTableRecord],
    pdf_path: Path,
    docai_dir: Path,
) -> None:
    payload = {
        "pdf_path": str(pdf_path),
        "docai_dir": str(docai_dir),
        "pdf_pages": [asdict(item) for item in pdf_pages],
        "docai_pages": [asdict(item) for item in docai_pages],
        "docai_tables": [asdict(item) for item in docai_tables],
    }
    manifest_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Exploit a degraded PDF with Document AI JSON + local OCR.")
    parser.add_argument("--pdf", default=str(PDF_PATH), help="Absolute path to the PDF file.")
    parser.add_argument("--docai-dir", default=str(DOCAI_JSON_DIR), help="Directory with Document AI JSON files.")
    parser.add_argument("--output-dir", default=str(OUTPUTS_DIR / "sin_titulo_extraction_full"), help="Directory where outputs will be written.")
    parser.add_argument("--tesseract-cmd", default=None, help="Optional explicit path to tesseract.exe.")
    parser.add_argument("--tessdata-dir", default=str(TESSDATA_DIR), help="Optional tessdata directory with spa/eng/osd.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    pdf_path = Path(args.pdf)
    docai_dir = Path(args.docai_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    tesseract_cmd = resolve_tesseract(args.tesseract_cmd)
    tessdata_dir = Path(args.tessdata_dir) if args.tessdata_dir else None
    configure_tesseract(tesseract_cmd=tesseract_cmd, tessdata_dir=tessdata_dir)

    docai_pages, docai_tables = load_docai_pages(docai_dir, output_dir)
    seed_docai_mappings(docai_pages)
    update_table_mappings(docai_pages, docai_tables)

    pdf_doc = fitz.open(str(pdf_path))
    covered_pages = {item.mapped_pdf_page for item in docai_pages if item.mapped_pdf_page is not None}
    ocr_targets = {
        page_number
        for page_number in range(1, pdf_doc.page_count + 1)
        if page_number not in covered_pages
    }
    ocr_targets.update({1, 2, 35, 49, 50, 65, 66, 67, 68, 69})

    ocr_pages: list[PdfPageOcr] = []
    token_rows: list[dict[str, Any]] = []
    for page_index in sorted(page_number - 1 for page_number in ocr_targets if 1 <= page_number <= pdf_doc.page_count):
        page = pdf_doc.load_page(page_index)
        ocr_result, rows = ocr_pdf_page(page)
        ocr_pages.append(ocr_result)
        token_rows.extend(rows)

    pdf_pages = synthesize_pdf_pages(
        page_count=pdf_doc.page_count,
        docai_pages=docai_pages,
        ocr_pages=ocr_pages,
    )

    write_text_exports(output_dir, pdf_pages, docai_pages)
    write_csv_exports(output_dir, pdf_pages, token_rows, docai_pages)
    write_excel(
        out_path=output_dir / "document_extraction.xlsx",
        pdf_pages=pdf_pages,
        docai_pages=docai_pages,
        docai_tables=docai_tables,
        table_csv_dir=output_dir / "tables_csv",
    )
    write_pdf_report(
        report_path=output_dir / "document_extraction_report.pdf",
        pdf_pages=pdf_pages,
        docai_pages=docai_pages,
        docai_tables=docai_tables,
        pdf_path=pdf_path,
        docai_dir=docai_dir,
    )
    write_docx_report(
        report_path=output_dir / "document_extraction_report.docx",
        pdf_pages=pdf_pages,
        docai_pages=docai_pages,
        docai_tables=docai_tables,
        pdf_path=pdf_path,
        docai_dir=docai_dir,
    )
    write_processing_manifest(
        manifest_path=output_dir / "processing_manifest.json",
        pdf_pages=pdf_pages,
        docai_pages=docai_pages,
        docai_tables=docai_tables,
        pdf_path=pdf_path,
        docai_dir=docai_dir,
    )

    print(f"Done. Outputs written to: {output_dir}")


if __name__ == "__main__":
    main()
