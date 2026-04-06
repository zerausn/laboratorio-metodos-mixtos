"""
backend/export_module.py
========================
Funciones de exportación multi-formato para el laboratorio de métodos mixtos.

Formatos soportados:
  - CSV  : tablas de datos, resultados OCR
  - Excel: resultados OCR, libros de códigos, matrices cualitativas
  - Word : informes narrativos de OCR, proyectos de codificación
  - PDF  : versión imprimible (requiere reportlab)

Todas las funciones aceptan también un bytes-buffer en vez de ruta en disco,
útil para Streamlit (st.download_button).
"""

from __future__ import annotations

import io
from pathlib import Path

import pandas as pd

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DocxDocument = None
    _DOCX_OK = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    _REPORTLAB_OK = True
except ImportError:
    _REPORTLAB_OK = False

try:
    import xlsxwriter  # noqa: F401 — validar disponibilidad
    _XLSXWRITER_OK = True
except ImportError:
    _XLSXWRITER_OK = False


# ── Helpers ───────────────────────────────────────────────────────────────────

def _prepare(output_path) -> Path:
    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


# ── Clase principal ───────────────────────────────────────────────────────────

class DataExporter:
    """
    Exportador multi-formato. Todos los métodos `to_*` devuelven la ruta
    del archivo generado (o None si falta dependencia).

    Los métodos `to_*_bytes` devuelven un BytesIO listo para Streamlit
    sin tocar el disco.
    """

    # ── CSV ───────────────────────────────────────────────────────────────────

    @staticmethod
    def to_csv(data_list, output_path) -> Path:
        output_path = _prepare(output_path)
        pd.DataFrame(data_list).to_csv(output_path, index=False, encoding="utf-8")
        return output_path

    @staticmethod
    def to_csv_bytes(data_list) -> bytes:
        return pd.DataFrame(data_list).to_csv(index=False, encoding="utf-8").encode("utf-8")

    # ── Excel ─────────────────────────────────────────────────────────────────

    @staticmethod
    def to_excel(data_list, output_path) -> Path | None:
        output_path = _prepare(output_path)
        engine = "xlsxwriter" if _XLSXWRITER_OK else "openpyxl"
        try:
            pd.DataFrame(data_list).to_excel(output_path, index=False, engine=engine)
        except Exception as exc:
            print(f"[export_module] Error Excel: {exc}")
            return None
        return output_path

    @staticmethod
    def to_excel_bytes(data_list) -> bytes | None:
        engine = "xlsxwriter" if _XLSXWRITER_OK else "openpyxl"
        buf = io.BytesIO()
        try:
            pd.DataFrame(data_list).to_excel(buf, index=False, engine=engine)
        except Exception as exc:
            print(f"[export_module] Error Excel bytes: {exc}")
            return None
        return buf.getvalue()

    # ── Excel con múltiples hojas (para resultados OCR detallados) ────────────

    @staticmethod
    def ocr_results_to_excel_bytes(ocr_results: list) -> bytes | None:
        """
        Exporta resultados OCR multi-engine a Excel con:
        - Hoja 'Resumen': página, mejor motor, score, texto breve
        - Hoja 'Texto_Completo': texto completo por página
        - Hoja 'Scores_Motores': tabla comparativa de scores por motor
        """
        engine = "xlsxwriter" if _XLSXWRITER_OK else "openpyxl"
        buf = io.BytesIO()
        try:
            rows_resumen = []
            rows_texto = []
            rows_scores = []
            for r in ocr_results:
                rows_resumen.append({
                    "Página": r.get("page"),
                    "Mejor Motor": r.get("best_engine"),
                    "Score": round(r.get("best_score", 0), 3),
                    "Combinado": r.get("combined", False),
                    "Preview (200 chars)": str(r.get("best_text", ""))[:200],
                })
                rows_texto.append({
                    "Página": r.get("page"),
                    "Texto Completo": r.get("best_text", ""),
                })
                score_row = {"Página": r.get("page")}
                score_row.update(r.get("score_summary", {}))
                rows_scores.append(score_row)

            with pd.ExcelWriter(buf, engine=engine) as writer:
                pd.DataFrame(rows_resumen).to_excel(writer, sheet_name="Resumen", index=False)
                pd.DataFrame(rows_texto).to_excel(writer, sheet_name="Texto_Completo", index=False)
                pd.DataFrame(rows_scores).to_excel(writer, sheet_name="Scores_Motores", index=False)
        except Exception as exc:
            print(f"[export_module] Error OCR Excel: {exc}")
            return None
        return buf.getvalue()

    # ── Codebook / Proyecto cualitativo a Excel ───────────────────────────────

    @staticmethod
    def codebook_to_excel_bytes(codebook: list, quotes_df: pd.DataFrame) -> bytes | None:
        """
        Exporta el libro de códigos y las citas codificadas a Excel.
        - Hoja 'Codebook': lista de códigos
        - Hoja 'Citas': tabla de fragmentos codificados
        - Hoja 'Matriz': tabla de co-ocurrencia código × documento
        """
        engine = "xlsxwriter" if _XLSXWRITER_OK else "openpyxl"
        buf = io.BytesIO()
        try:
            with pd.ExcelWriter(buf, engine=engine) as writer:
                pd.DataFrame({"Código": codebook}).to_excel(writer, sheet_name="Codebook", index=False)
                if not quotes_df.empty:
                    quotes_df.to_excel(writer, sheet_name="Citas", index=False)
                    try:
                        matriz = pd.crosstab(quotes_df["Documento"], quotes_df["Codigo"])
                        matriz.to_excel(writer, sheet_name="Matriz")
                    except Exception:
                        pass
        except Exception as exc:
            print(f"[export_module] Error Codebook Excel: {exc}")
            return None
        return buf.getvalue()

    # ── Word ──────────────────────────────────────────────────────────────────

    @staticmethod
    def to_word(data_list, output_path, title: str = "Resultado de extracción OCR") -> Path | None:
        if not _DOCX_OK:
            print("[export_module] python-docx no instalado.")
            return None
        output_path = _prepare(output_path)
        doc = _DocxDocument()
        doc.add_heading(title, 0)
        for item in data_list:
            page = item.get("page", "N/A")
            text = item.get("best_text") or item.get("text", "")
            engine = item.get("best_engine", "")
            score = item.get("best_score", item.get("score", ""))
            doc.add_heading(f"Página {page}", level=1)
            if engine:
                doc.add_paragraph(f"Motor: {engine} | Score: {score:.3f}" if isinstance(score, float) else f"Motor: {engine}")
            doc.add_paragraph(str(text))
            doc.add_page_break()
        doc.save(output_path)
        return output_path

    @staticmethod
    def to_word_bytes(data_list, title: str = "Resultado de extracción OCR") -> bytes | None:
        if not _DOCX_OK:
            return None
        doc = _DocxDocument()
        doc.add_heading(title, 0)
        for item in data_list:
            page = item.get("page", "N/A")
            text = item.get("best_text") or item.get("text", "")
            engine = item.get("best_engine", "")
            score = item.get("best_score", item.get("score", ""))
            doc.add_heading(f"Página {page}", level=1)
            if engine:
                try:
                    doc.add_paragraph(f"Motor: {engine} | Score: {score:.3f}")
                except Exception:
                    doc.add_paragraph(f"Motor: {engine}")
            doc.add_paragraph(str(text))
            doc.add_page_break()
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── PDF ───────────────────────────────────────────────────────────────────

    @staticmethod
    def to_pdf_text(data_list, output_path) -> Path | None:
        if not _REPORTLAB_OK:
            print("[export_module] reportlab no instalado.")
            return None
        output_path = _prepare(output_path)
        pdf = canvas.Canvas(str(output_path), pagesize=letter)
        _, height = letter
        for item in data_list:
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(50, height - 50, f"Página {item.get('page', 'N/A')}")
            text_object = pdf.beginText(50, height - 80)
            text_object.setFont("Helvetica", 9)
            for line in str(item.get("best_text") or item.get("text", "")).splitlines():
                text_object.textLine(line[:110])
            pdf.drawText(text_object)
            pdf.showPage()
        pdf.save()
        return output_path
