from pathlib import Path

import pandas as pd
from docx import Document

try:
    from backend.path_utils import reports_dir
except ImportError:
    from path_utils import reports_dir


SAMPLE_DATA = [
    {
        "Rubro": "Gastos de Funcionamiento",
        "Ppto. Modificado": 1133833863257,
        "Pagos": 800734794257,
        "Saldo Obligaciones": 12461335150,
        "Ejecucion (%)": 77.98,
    },
    {
        "Rubro": "Inversion",
        "Ppto. Modificado": 5861000000000,
        "Pagos": 4014000000000,
        "Saldo Obligaciones": 0,
        "Ejecucion (%)": 68.49,
    },
    {
        "Rubro": "GRAN TOTAL DISTRITO",
        "Ppto. Modificado": 7188000000000,
        "Pagos": 5036000000000,
        "Saldo Obligaciones": 0,
        "Ejecucion (%)": 70.06,
    },
]


def generate_reports(data=None, output_dir=None):
    data = data or SAMPLE_DATA
    output_dir = Path(output_dir or reports_dir())
    output_dir.mkdir(parents=True, exist_ok=True)

    dataframe = pd.DataFrame(data)
    excel_path = output_dir / "TABLA_PRESUPUESTO_RECUPERADA.xlsx"
    dataframe.to_excel(excel_path, index=False, engine="xlsxwriter")

    doc = Document()
    doc.add_heading("Informe de ejecucion presupuestal - Santiago de Cali", 0)
    doc.add_paragraph("Datos recuperados mediante OCR y validacion local.")
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    for index, column in enumerate(dataframe.columns):
        table.rows[0].cells[index].text = str(column)

    for row in dataframe.itertuples(index=False):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)

    word_path = output_dir / "INFORME_EJECUCION_CALI.docx"
    doc.save(word_path)
    return {"excel": excel_path, "word": word_path}


if __name__ == "__main__":
    generate_reports()
