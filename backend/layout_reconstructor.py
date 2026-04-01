import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

try:
    from backend.path_utils import default_checkpoint_path, reports_dir
except ImportError:
    from path_utils import default_checkpoint_path, reports_dir


def _extract_rows(text):
    rows = []
    for line in str(text).splitlines():
        stripped = line.strip()
        if len(stripped) <= 5:
            continue
        cols = [part for part in re.split(r"\s{3,}|\|", stripped) if part]
        if len(cols) > 1:
            rows.append(cols)
    return rows


def reconstruct_layout(json_data_path=None, output_dir=None):
    json_data_path = Path(json_data_path or default_checkpoint_path())
    output_dir = Path(output_dir or reports_dir())
    if not json_data_path.exists():
        return None

    output_dir.mkdir(parents=True, exist_ok=True)
    with open(json_data_path, "r", encoding="utf-8") as handle:
        status = json.load(handle)

    excel_path = output_dir / "RECONSTRUCCION_TOTAL_CALI.xlsx"
    writer = pd.ExcelWriter(excel_path, engine="xlsxwriter")
    doc = Document()
    doc.add_heading("Reconstruccion de ejecucion presupuestal Cali - Sep 2025", 0)
    pdf_path = output_dir / "RECONSTRUCCION_TOTAL_CALI.pdf"
    pdf = canvas.Canvas(str(pdf_path), pagesize=letter)
    _, height = letter

    for entry in status.get("data", []):
        page_num = entry.get("page", "N/A")
        text = entry.get("tesseract") or entry.get("text", "")
        rows = _extract_rows(text)
        if not rows:
            continue

        max_cols = max(len(row) for row in rows)
        normalized = [row + [""] * (max_cols - len(row)) for row in rows]
        dataframe = pd.DataFrame(normalized)
        dataframe.to_excel(writer, sheet_name=f"Pag {page_num}", index=False)

        doc.add_heading(f"Pagina {page_num}", level=1)
        table = doc.add_table(rows=len(normalized), cols=max_cols)
        for row_index, rowdata in enumerate(normalized):
            for col_index, value in enumerate(rowdata):
                table.cell(row_index, col_index).text = str(value)
        doc.add_page_break()

        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(40, height - 40, f"Pagina {page_num}")
        text_object = pdf.beginText(40, height - 65)
        text_object.setFont("Helvetica", 8)
        for rowdata in normalized:
            text_object.textLine(" | ".join(str(value) for value in rowdata))
        pdf.drawText(text_object)
        pdf.showPage()

    writer.close()
    word_path = output_dir / "RECONSTRUCCION_TOTAL_CALI.docx"
    doc.save(word_path)
    pdf.save()
    return {"excel": excel_path, "word": word_path, "pdf": pdf_path}


if __name__ == "__main__":
    reconstruct_layout()
